"""
DOCX extractor for DOCINTEL.
Extracts text and metadata from DOCX files using python-docx.
"""

from pathlib import Path
from typing import Optional

import docx

from docintel.common.exceptions import ExtractionError
from docintel.common.logging import get_logger
from docintel.common.constants import SOURCE_LAYER_NATIVE

logger = get_logger(__name__)


class DOCXExtractor:
    """
    Extracts text and metadata from DOCX files.
    Uses python-docx for reading Word documents.
    """
    
    def __init__(self):
        """Initialize the DOCX extractor."""
        pass
    
    def extract(self, file_path: str) -> dict:
        """
        Extract text and metadata from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary with 'text', 'metadata', and 'source_layer'
        """
        try:
            doc = docx.Document(file_path)
            
            # Extract metadata
            metadata = self._extract_metadata(doc)
            
            # Extract text from all paragraphs
            text_parts = []
            
            # Extract from main document
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract from tables
            for table_num, table in enumerate(doc.tables, 1):
                table_text = self._extract_table_text(table, table_num)
                if table_text:
                    text_parts.append(f"\n[Table {table_num}]\n{table_text}")
            
            full_text = "\n".join(text_parts)
            
            logger.info(
                "DOCX extraction complete",
                extra={
                    "file": file_path,
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables),
                    "total_chars": len(full_text),
                },
            )
            
            return {
                "text": full_text,
                "metadata": metadata,
                "source_layer": SOURCE_LAYER_NATIVE,
            }
            
        except Exception as e:
            error_msg = f"Failed to extract DOCX: {str(e)}"
            logger.error(error_msg, extra={"file": file_path}, exc_info=True)
            raise ExtractionError(error_msg) from e
    
    def _extract_table_text(self, table, table_num: int) -> str:
        """
        Extract text from a DOCX table.
        
        Args:
            table: python-docx table object
            table_num: Table number for logging
            
        Returns:
            Extracted table text
        """
        try:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            return "\n".join(rows)
        except Exception as e:
            logger.warning(f"Could not extract table {table_num}: {str(e)}")
            return ""
    
    def _extract_metadata(self, doc: docx.Document) -> dict:
        """
        Extract metadata from a DOCX document.
        
        Args:
            doc: python-docx document object
            
        Returns:
            Dictionary of metadata
        """
        metadata = {}
        
        try:
            # Get core properties
            core_props = doc.core_properties
            metadata["title"] = core_props.title or ""
            metadata["author"] = core_props.author or ""
            metadata["subject"] = core_props.subject or ""
            metadata["created"] = str(core_props.created) if core_props.created else ""
            metadata["modified"] = str(core_props.modified) if core_props.modified else ""
            metadata["last_modified_by"] = core_props.last_modified_by or ""
            
            # Count elements
            metadata["paragraph_count"] = len(doc.paragraphs)
            metadata["table_count"] = len(doc.tables)
            
        except Exception as e:
            logger.warning(f"Could not extract all DOCX metadata: {str(e)}")
        
        return metadata
    
    def extract_page(self, file_path: str, page_number: int) -> dict:
        """
        DOCX files don't have pages in the same way as PDFs.
        Returns the full extraction for consistency.
        
        Args:
            file_path: Path to the DOCX file
            page_number: Ignored for DOCX
            
        Returns:
            Dictionary with 'text' and 'metadata'
        """
        return self.extract(file_path)
    
    def get_paragraph_count(self, file_path: str) -> int:
        """
        Get the number of paragraphs in a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Number of paragraphs
        """
        try:
            doc = docx.Document(file_path)
            return len(doc.paragraphs)
        except Exception as e:
            error_msg = f"Failed to count paragraphs: {str(e)}"
            logger.error(error_msg, extra={"file": file_path})
            raise ExtractionError(error_msg) from e


# Global DOCX extractor instance
_docx_extractor: Optional[DOCXExtractor] = None


def get_docx_extractor() -> DOCXExtractor:
    """Get the global DOCX extractor instance."""
    global _docx_extractor
    if _docx_extractor is None:
        _docx_extractor = DOCXExtractor()
    return _docx_extractor